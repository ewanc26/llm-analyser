"""Document reading utilities for .docx files."""

from pathlib import Path
from typing import Dict

from docx import Document


def read_docx(file_path: Path) -> Dict:
    """Extract paragraphs and tables from a .docx file."""
    doc = Document(file_path)

    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    tables_content = []
    for table in doc.tables:
        table_text = []
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                table_text.append(" | ".join(row_text))
        if table_text:
            tables_content.append("\n".join(table_text))

    return {
        "paragraphs": "\n\n".join(paragraphs),
        "tables": "\n\n".join(tables_content) if tables_content else "",
        "word_count": len(" ".join(paragraphs).split()),
        "paragraph_count": len(paragraphs),
    }
